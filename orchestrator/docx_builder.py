"""
================================================================================
 I-SCEET — DOCX Builder v2.1
 File: orchestrator/docx_builder.py
================================================================================
 Parses M1/M2 LLM text output -> generates professional Word .docx
 M1 output -> SRD-0xx.docx  (Sections I, II, III.1, III.2, IV)
 M2 output -> SDDD-0xx.docx (Section 1: Architecture, Section 2: LLRs)

 Updated for M1_system_prompt v3.1 enriched HLR format:
   ID, Title, Activation Period, Statement, Input Data, Output Data,
   Local Data, Rationale, Parent, Upwards Traceability, Terminal,
   Category, Context, Priority, Status, Verification Method,
   Safety Effect, Security Assurance Level
================================================================================
"""

import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError("python-docx required: pip install python-docx")

# ── STYLE CONSTANTS ───────────────────────────────────────────────────────────
BLUE     = RGBColor(0x1F, 0x38, 0x64)
MID_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GRAY     = RGBColor(0x66, 0x66, 0x66)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "Arial"

# ── Precise HLR block colors (per approved style spec) ─────────────────────────
ID_BLUE      = RGBColor(0x3B, 0x82, 0xF6)  # HLR-xxx heading, [END_xxx] marker
MAROON       = RGBColor(0x77, 0x00, 0x29)  # <Activation Period: ...>
DARK_TITLE   = RGBColor(0x1C, 0x1C, 0x1C)  # "IV.1.1 - Capability Title" heading
BLACK        = RGBColor(0x00, 0x00, 0x00)
SERIF_FONT   = "Liberation Serif"          # metadata field list font

# HLR field order — used both for boundary-aware parsing and rendering
HLR_FIELDS = [
    'Title', 'Activation Period', 'Statement', 'Input Data', 'Output Data',
    'Local Data', 'Rationale', 'Parent', 'Upwards Traceability', 'Terminal',
    'Category', 'Context', 'Priority', 'Status', 'Verification Method',
    'Safety Effect', 'Security Assurance Level',
]


# ── TEXT CLEANUP ──────────────────────────────────────────────────────────────
def _strip_think_tags(text: str) -> str:
    """Remove <think>...</think> reasoning artifacts (defense in depth,
    also stripped server-side in colab_server.py)."""
    if not text:
        return text
    cleaned = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    cleaned = re.sub(r'</?think>', '', cleaned)
    return cleaned.strip()


# ── DOCX HELPERS ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '2E75B6')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = MID_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.color.rgb = MID_BLUE
        run.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
    return p


def _add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    return p


def _add_header_footer(doc, doc_id, version, date_str):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    run = hp.add_run(f"{doc_id}   |   I-SCEET   |   v{version}   |   DRAFT")
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), '2E75B6')
    pBdr.append(bot)
    pPr.append(pBdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    run2 = fp.add_run(f"CONFIDENTIAL — Capgemini Engineering / ESPRIT   |   {date_str}")
    run2.font.name = FONT_NAME
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    pPr2 = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:color'), '2E75B6')
    pBdr2.append(top)
    pPr2.append(pBdr2)


def _make_cover(doc, doc_id, title, version, project, author, date_str, dal):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = BLUE
    r.font.name = FONT_NAME
    p_title.paragraph_format.space_before = Pt(60)
    p_title.paragraph_format.space_after  = Pt(8)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Intelligent Safety-Critical Engineering Environment Toolchain")
    r2.font.size = Pt(14)
    r2.font.color.rgb = GRAY
    r2.font.name = FONT_NAME
    r2.italic = True
    p_sub.paragraph_format.space_after = Pt(30)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Document ID",  doc_id),
        ("Version",      version),
        ("Status",       "DRAFT"),
        ("Date",         date_str),
        ("Author",       author),
        ("Project",      project),
        ("Institution",  "Capgemini Engineering / ESPRIT"),
        ("DAL Level",    dal),
        ("Standard",     "RTCA DO-178C §5"),
    ]
    for i, (key, val) in enumerate(rows_data):
        row = table.add_row()
        row.cells[0].text = key
        row.cells[1].text = val
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(10)
        _set_cell_bg(row.cells[0], "1F3864" if i == 0 else "F2F2F2")
        run_k = row.cells[0].paragraphs[0].runs[0]
        run_k.bold = True
        run_k.font.name = FONT_NAME
        run_k.font.size = Pt(10)
        run_k.font.color.rgb = WHITE if i == 0 else MID_BLUE
        run_v = row.cells[1].paragraphs[0].runs[0]
        run_v.font.name = FONT_NAME
        run_v.font.size = Pt(10)

    doc.add_page_break()


def _make_revision_table(doc, version, date_str, generator="M1 (I-SCEET AI)",
                          revision_history: list = None):
    """
    Render the Document Revision History table.
    revision_history: optional list of past entries (dicts with keys
    version/date/author/description) rendered BEFORE the current version row.
    """
    _add_heading(doc, "Document Revision History")
    table = doc.add_table(rows=1, cols=4)
    headers = ["Version", "Date", "Author", "Description"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "1F3864")
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.name = FONT_NAME
        r.font.size = Pt(10)

    all_rows = []
    if revision_history:
        for entry in revision_history:
            all_rows.append([
                str(entry.get('version', '')),
                str(entry.get('date', '')),
                str(entry.get('author', generator)),
                str(entry.get('description', 'AI-generated document')),
            ])
    all_rows.append([str(version), date_str, generator, "AI-generated document"])

    for vals in all_rows:
        row = table.add_row()
        for j, v in enumerate(vals):
            row.cells[j].text = v
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            row.cells[j].paragraphs[0].runs[0].font.name = FONT_NAME

    doc.add_paragraph()


def _add_pipe_table(doc, lines: list):
    """Render markdown-style '| a | b | c |' lines as a real Word table."""
    rows = []
    for l in lines:
        l = l.strip()
        if not l.startswith('|'):
            continue
        cells = [c.strip() for c in l.strip('|').split('|')]
        if all(re.fullmatch(r'-+', c) for c in cells):
            continue  # separator row
        rows.append(cells)
    if not rows:
        return False

    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncols)
    for i, row_vals in enumerate(rows):
        row = tbl.add_row()
        for j in range(ncols):
            text = row_vals[j] if j < len(row_vals) else ""
            row.cells[j].text = text
            if row.cells[j].paragraphs[0].runs:
                run = row.cells[j].paragraphs[0].runs[0]
                run.font.name = FONT_NAME
                run.font.size = Pt(9.5)
                if i == 0:
                    run.bold = True
                    run.font.color.rgb = WHITE
            if i == 0:
                _set_cell_bg(row.cells[j], "1F3864")
            elif i % 2 == 0:
                _set_cell_bg(row.cells[j], "F2F2F2")
    doc.add_paragraph()
    return True


def _render_multiline_field(doc, content: str):
    """Render a field's content: as a table if it contains pipe-table lines,
    otherwise as bullet points / plain paragraphs."""
    if not content or not content.strip():
        _add_body(doc, "None")
        return
    lines = [l for l in content.split('\n') if l.strip()]
    pipe_lines = [l for l in lines if l.strip().startswith('|')]
    if len(pipe_lines) >= 2:
        _add_pipe_table(doc, lines)
        # Render any non-table lines too (e.g. intro sentence before table)
        remaining = [l for l in lines if not l.strip().startswith('|')]
        for l in remaining:
            _add_body(doc, l.strip())
    else:
        for l in lines:
            l = l.strip()
            if l.startswith('-') or l.startswith('•'):
                _add_bullet(doc, l.lstrip('-• '))
            else:
                _add_body(doc, l)


# ── REQUIREMENT CARD STYLE HELPERS (industrial SRD look) ──────────────────────
def _remove_cell_borders(cell):
    """Strip all borders from a table cell (used for borderless field lists)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tcPr.append(borders)


def _add_id_heading(doc, id_text: str):
    """Requirement ID heading (e.g. HLR-BOOT-001).
    Arial 11pt, bold italic, color #3B82F6 (per approved style spec)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(id_text)
    run.bold = True
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = ID_BLUE
    return p


def _add_italic_blue(doc, text: str, size: float = 9.5, color=None):
    """Small italic line — color defaults to MID_BLUE, override for
    Activation Period (MAROON) or END marker (ID_BLUE)."""
    if color is None:
        color = MID_BLUE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(text)
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _add_bold_italic_title(doc, text: str):
    """Capability title line — bold italic black text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = FONT_NAME
    run.font.size = Pt(11.5)
    return p


def _add_horizontal_rule(doc):
    """Thin horizontal separator line — minimal height (tiny font on the
    empty run prevents the paragraph from taking a full default line height)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("")
    run.font.size = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_subsystem_heading(doc, text: str):
    """Subsystem group heading, e.g. 'IV.1 — BOOT Requirements'.
    Arial 16pt, bold, color #2E75B6 (per approved style spec)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.color.rgb = MID_BLUE
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    return p


def _add_numbered_heading(doc, text: str):
    """Numbered requirement-title heading, e.g. 'IV.1.1 - Initialize hardware...'
    Arial 13pt, bold, color #1C1C1C, NOT italic (per approved style spec)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_TITLE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p


def _add_numbered_subheading(doc, text: str):
    """Numbered sub-section heading, e.g. 'IV.1.1.1 Input Data'.
    Bold italic MID_BLUE, matches existing level=3 style."""
    return _add_heading(doc, text, 3)


def _add_field_line(doc, label: str, value: str, tab_pos_cm: float = 6.5):
    """
    Render one metadata field as a SINGLE paragraph line with a tab stop
    between label and value — NOT a table (per approved style spec).
    Uses a hanging indent so that if the value wraps to a second line,
    the continuation stays aligned in the value column instead of
    falling back to the left margin.
    """
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_pos_cm))
    p.paragraph_format.left_indent      = Cm(tab_pos_cm)
    p.paragraph_format.first_line_indent = Cm(-tab_pos_cm)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"{label}\t{value}")
    run.italic = True
    run.font.name = SERIF_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    return p


def _add_borderless_field_list(doc, pairs: list):
    """
    Render a label/value list as individual tab-aligned paragraph lines —
    NO table structure at all (per approved style spec).
    """
    for label, value in pairs:
        if not value:
            continue
        _add_field_line(doc, label, value)


# ── PARSER (v3.1 enriched HLR format) ──────────────────────────────────────────
def _clean_field_value(value: str) -> str:
    """
    Safety net: strip any leaked block-boundary artifacts from a field's
    captured value (closing '---' delimiters, [END_HLR-xxx] markers, the
    next subsystem header, or the start of the next HLR/LLR block). This
    protects against block-splitting edge cases where the LLM output format
    varies slightly (e.g. indentation before 'ID:').
    """
    if not value:
        return value
    cut_patterns = [
        r'\n\s*-{3,}',            # closing/opening --- delimiter
        r'\[END_HLR-',            # HLR end marker
        r'\[END_LLR-',            # LLR end marker
        r'\n\s*ID:\s*HLR-',       # start of a leaked next HLR block
        r'\n\s*ID:\s*LLR-',       # start of a leaked next LLR block
    ]
    earliest = len(value)
    for pat in cut_patterns:
        m = re.search(pat, value)
        if m and m.start() < earliest:
            earliest = m.start()
    return value[:earliest].strip()


def _parse_req_block(block: str) -> dict:
    """Parse one HLR/LLR block's fields using boundary-aware regex."""
    req = {}
    boundary = '(?:' + '|'.join(re.escape(f) + ':' for f in HLR_FIELDS) + ')'
    for i, field in enumerate(HLR_FIELDS):
        pattern = rf'{re.escape(field)}:\s*(.*?)(?=\n{boundary}|\Z)'
        m = re.search(pattern, block, re.DOTALL)
        raw_value = m.group(1).strip() if m else ''
        req[field] = _clean_field_value(raw_value)
    return req


def _parse_m1_output(text: str) -> dict:
    """Parse M1 output (v3.1 format) into structured sections."""
    text = _strip_think_tags(text)
    result = {
        'sec1': '', 'sec2': '', 'sec3_1': '',
        'sec3_2': '', 'sec4': '', 'hlrs': [], 'coverage': ''
    }

    sec_pattern = re.compile(
        r'[═=]{10,}\s*\n(SECTION [IVX]+ [—-][^\n]+)\n[═=]{10,}', re.MULTILINE
    )
    parts = sec_pattern.split(text)

    current = None
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if re.match(r'SECTION I\s*[—-]', part):
            current = 'sec1'
        elif re.match(r'SECTION II\s*[—-]', part):
            current = 'sec2'
        elif re.match(r'SECTION III\s*[—-]', part):
            current = 'sec3'
        elif re.match(r'SECTION IV\s*[—-]', part):
            current = 'sec4'
        else:
            if current == 'sec1':
                result['sec1'] = part
            elif current == 'sec2':
                result['sec2'] = part
            elif current == 'sec3':
                sub = re.split(r'III\.2\s+Dynamic', part, maxsplit=1, flags=re.IGNORECASE)
                result['sec3_1'] = sub[0].strip()
                result['sec3_2'] = sub[1].strip() if len(sub) > 1 else ''
            elif current == 'sec4':
                # Split off PDS Coverage Summary if present
                cov_split = re.split(r'PDS COVERAGE SUMMARY', part, maxsplit=1)
                hlr_text = cov_split[0]
                result['coverage'] = cov_split[1].strip() if len(cov_split) > 1 else ''
                result['sec4'] = hlr_text

                # Find each HLR block starting at "ID:" lines
                id_positions = [m.start() for m in re.finditer(r'(?m)^\s*ID:\s*HLR-', hlr_text)]
                for idx, pos in enumerate(id_positions):
                    end = id_positions[idx + 1] if idx + 1 < len(id_positions) else len(hlr_text)
                    block = hlr_text[pos:end]
                    id_match = re.match(r'\s*ID:\s*(HLR-[A-Z]+-\d+[A-Za-z-]*)', block)
                    if not id_match:
                        continue
                    hlr = {'ID': id_match.group(1)}
                    hlr.update(_parse_req_block(block))
                    result['hlrs'].append(hlr)

    return result


def _parse_m2_output(text: str) -> dict:
    text = _strip_think_tags(text)
    result = {'sec1': '', 'sec2': '', 'llrs': []}
    sec_pattern = re.compile(r'SECTION\s+([12])[:\s—-]+(.+?)(?=SECTION\s+[12][:\s—-]|\Z)',
                              re.DOTALL | re.IGNORECASE)
    for m in sec_pattern.finditer(text):
        num, content = m.group(1), m.group(2).strip()
        if num == '1':
            result['sec1'] = content
        elif num == '2':
            result['sec2'] = content
            id_positions = [mm.start() for mm in re.finditer(r'(?m)^\s*ID:\s*LLR-', content)]
            for idx, pos in enumerate(id_positions):
                end = id_positions[idx + 1] if idx + 1 < len(id_positions) else len(content)
                block = content[pos:end]
                id_match = re.match(r'\s*ID:\s*(LLR-[A-Z]+-\d+[A-Za-z-]*)', block)
                if not id_match:
                    continue
                llr = {'ID': id_match.group(1)}
                llr.update(_parse_req_block(block))
                result['llrs'].append(llr)
    return result


# ── HLR / LLR RENDERING ────────────────────────────────────────────────────────
def _add_req_card(doc, req: dict, prefix: str = ""):
    """
    Render one HLR/LLR matching the reference industrial SRD style exactly:

      <prefix> - Capability Title          (numbered heading, hyphen, bold)
      <Activation Period: ...>              (italic blue)
      <prefix>.1 Input Data                 (numbered sub-heading + table)
      <prefix>.2 Output Data                (numbered sub-heading + table)
      <prefix>.3 Local Data                 (numbered sub-heading + table/None)
      <prefix>.4 Requirement                (numbered sub-heading)
        HLR-xxx-NNN                         (bold, underlined)
        [COV.Parent]                        (italic blue)
        ─────────────────────────────
        Capability Title                    (bold italic)
        Statement text                      (plain paragraph)

        Upwards Traceability   ...
        Category                ...
        Context                 ...
        Verification Method     ...
        Security Assurance Level ...
        Safety Effect            ...
        Rationale                ...        (borderless field list)
        ─────────────────────────────
        [END_HLR-xxx-NNN]                   (italic blue)

    `prefix` is the hierarchical number, e.g. "IV.1.1" (no trailing dot).
    If prefix is empty, falls back to unnumbered rendering (used for M2/LLR
    reuse where a different numbering scheme may apply).
    """
    title_text = req.get('Title', '').strip('"')
    activation = req.get('Activation Period', '')

    # ── Numbered capability heading ────────────────────────────────────────
    heading_text = f"{prefix} - {title_text}" if prefix else (title_text or req['ID'])
    _add_numbered_heading(doc, heading_text)
    if activation:
        _add_italic_blue(doc, f"<Activation Period: {activation}>", size=10, color=MAROON)

    # ── Input / Output / Local Data — numbered, BEFORE the requirement ────
    data_specs = [('Input Data', 'Input Data'),
                  ('Output Data', 'Output Data'),
                  ('Local Data', 'Local Data')]
    for i, (data_field, label) in enumerate(data_specs, 1):
        val = req.get(data_field, '')
        sub_prefix = f"{prefix}.{i} " if prefix else ""
        _add_numbered_subheading(doc, f"{sub_prefix}{label}")
        if val and val.strip().lower() not in ('none', 'n/a', ''):
            _render_multiline_field(doc, val)
        else:
            _add_body(doc, "None")

    # ── Requirement sub-section ────────────────────────────────────────────
    req_prefix = f"{prefix}.4 " if prefix else ""
    _add_numbered_subheading(doc, f"{req_prefix}Requirement")

    _add_id_heading(doc, req['ID'])

    parent = req.get('Parent', '')
    if parent:
        _add_italic_blue(doc, f"[COV.{parent}]")
    _add_horizontal_rule(doc)

    if title_text:
        _add_bold_italic_title(doc, title_text)

    statement = req.get('Statement', '')
    if statement:
        _add_body(doc, statement)

    # ── Metadata — borderless field list (no "Terminal", matches reference) ─
    meta_pairs = [
        ("Upwards Traceability", req.get('Upwards Traceability', '')),
        ("Category",             req.get('Category', '')),
        ("Context",              req.get('Context', '')),
        ("Verification Method",  req.get('Verification Method', '')),
        ("Security Assurance Level", req.get('Security Assurance Level', '')),
        ("Safety Effect",        req.get('Safety Effect', '')),
        ("Rationale",            req.get('Rationale', '')),
    ]
    _add_borderless_field_list(doc, meta_pairs)

    _add_horizontal_rule(doc)
    _add_italic_blue(doc, f"[END_{req['ID']}]", color=ID_BLUE)
    doc.add_paragraph()


def _extract_doc_id(sec1_text: str, prefix: str, version: str) -> str:
    """
    Determine the real document ID:
    1. Prefer what M1 itself declared in Section I.3 ("Document ID: SRD-002")
    2. Fallback: derive from the version number (e.g. version=3 -> SRD-003)
    """
    if sec1_text:
        m = re.search(rf'Document ID:\s*({prefix}-[A-Za-z0-9]+)', sec1_text)
        if m:
            return m.group(1)
    try:
        num = int(float(version))
        return f"{prefix}-{num:03d}"
    except (ValueError, TypeError):
        return f"{prefix}-001"


# ── MAIN BUILDERS ─────────────────────────────────────────────────────────────
def build_srd(m1_output: str, output_path: str, version: str = "1.0",
              project: str = "I-SCEET", revision_history: list = None) -> str:
    """Build SRD-xxx.docx from M1 v3.1 text output.

    revision_history: optional list of dicts (previous versions) each with
    keys 'version', 'date', 'author', 'description'. If provided, all past
    entries are shown in the Document Revision History table, followed by
    the current generation as the last row.
    """
    date_str = datetime.now().strftime("%Y-%m-%d")

    # Parse first so we can determine the real doc_id before building the cover
    parsed = _parse_m1_output(m1_output)
    doc_id = _extract_doc_id(parsed['sec1'], 'SRD', version)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)

    _make_cover(doc, doc_id, "SOFTWARE REQUIREMENTS DOCUMENT",
                version, project, "M1 (I-SCEET — Qwen3 32B)", date_str, "A")
    _add_header_footer(doc, doc_id, version, date_str)
    _make_revision_table(doc, version, date_str,
                        generator="M1 (I-SCEET AI)",
                        revision_history=revision_history)
    doc.add_page_break()

    # Section I
    _add_heading(doc, "SECTION I — INTRODUCTION", 1)
    if parsed['sec1']:
        for line in parsed['sec1'].split('\n'):
            line = line.strip()
            if not line:
                continue
            if re.match(r'I\.\d\s', line):
                _add_heading(doc, line, 2)
            else:
                _add_body(doc, line)
    else:
        _add_body(doc, "[Section I not found in model output]")
    doc.add_page_break()

    # Section II
    _add_heading(doc, "SECTION II — REFERENCED DOCUMENTS", 1)
    if parsed['sec2']:
        lines = [l for l in parsed['sec2'].split('\n') if l.strip()]
        if not _add_pipe_table(doc, lines):
            for l in lines:
                _add_body(doc, l.strip())
    doc.add_page_break()

    # Section III
    _add_heading(doc, "SECTION III — SYSTEM ARCHITECTURE", 1)
    _add_heading(doc, "III.1 Static Architecture", 2)
    if parsed['sec3_1']:
        _render_multiline_field(doc, parsed['sec3_1'])
    else:
        _add_body(doc, "[Static architecture not found in model output]")

    _add_heading(doc, "III.2 Dynamic Architecture", 2)
    if parsed['sec3_2']:
        _render_multiline_field(doc, parsed['sec3_2'])
    else:
        _add_body(doc, "[Dynamic architecture not found in model output]")
    doc.add_page_break()

    # Section IV
    _add_heading(doc, "IV — HIGH LEVEL REQUIREMENTS", 1)
    if parsed['hlrs']:
        by_subsystem = {}
        for hlr in parsed['hlrs']:
            sub = re.search(r'HLR-([A-Z]+)-', hlr.get('ID', ''))
            key = sub.group(1) if sub else 'OTHER'
            by_subsystem.setdefault(key, []).append(hlr)

        for subsystem_idx, (subsystem, hlrs) in enumerate(by_subsystem.items(), 1):
            _add_subsystem_heading(doc, f"IV.{subsystem_idx} — {subsystem} Requirements")
            for req_idx, hlr in enumerate(hlrs, 1):
                _add_req_card(doc, hlr, prefix=f"IV.{subsystem_idx}.{req_idx}")

        _add_heading(doc, "IV.SUMMARY — PDS Coverage", 2)
        derived = sum(1 for h in parsed['hlrs'] if '-D' in h.get('ID', ''))
        _add_body(doc, f"Total HLRs generated: {len(parsed['hlrs'])}")
        _add_body(doc, f"Standard HLRs: {len(parsed['hlrs']) - derived}")
        _add_body(doc, f"Derived HLRs (-D): {derived}")
        if parsed['coverage']:
            doc.add_paragraph()
            _render_multiline_field(doc, parsed['coverage'])
    else:
        _add_body(doc, "[No HLRs could be parsed from model output — raw text below]")
        for line in (parsed['sec4'] or m1_output).split('\n')[:150]:
            _add_body(doc, line.strip())

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    return output_path


def build_sddd(m2_output: str, output_path: str, version: str = "1.0",
               project: str = "I-SCEET") -> str:
    """Build SDDD-0xx.docx from M2 text output."""
    date_str = datetime.now().strftime("%Y-%m-%d")
    doc_id   = "SDDD-0xx"

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)

    _make_cover(doc, doc_id, "SOFTWARE DESIGN DESCRIPTION DOCUMENT",
                version, project, "M2 (I-SCEET — Mistral Small 24B)", date_str, "A")
    _add_header_footer(doc, doc_id, version, date_str)
    _make_revision_table(doc, version, date_str, generator="M2 (I-SCEET AI)")
    doc.add_page_break()

    parsed = _parse_m2_output(m2_output)

    _add_heading(doc, "SECTION 1 — SOFTWARE ARCHITECTURE", 1)
    if parsed['sec1']:
        _render_multiline_field(doc, parsed['sec1'])
    doc.add_page_break()

    _add_heading(doc, "SECTION 2 — LOW LEVEL REQUIREMENTS", 1)
    if parsed['llrs']:
        by_subsystem = {}
        for llr in parsed['llrs']:
            sub = re.search(r'LLR-([A-Z]+)-', llr.get('ID', ''))
            key = sub.group(1) if sub else 'OTHER'
            by_subsystem.setdefault(key, []).append(llr)
        for subsystem_idx, (subsystem, llrs) in enumerate(by_subsystem.items(), 1):
            _add_heading(doc, f"2.{subsystem_idx} — {subsystem} Low Level Requirements", 2)
            for req_idx, llr in enumerate(llrs, 1):
                _add_req_card(doc, llr, prefix=f"2.{subsystem_idx}.{req_idx}")
        _add_body(doc, f"Total LLRs generated: {len(parsed['llrs'])}")
    else:
        for line in (parsed['sec2'] or '').split('\n')[:200]:
            _add_body(doc, line.strip())

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    return output_path